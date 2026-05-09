"""Desirable/Undesirable sample-clause docx importer.

Targets the "Sample AI Playbook" format:
  Table columns: Topic | Sources | AI Playbook Rule | Desirable Sample Clause | Undesirable Sample Clause

One rule is created per non-header data row where 'AI Playbook Rule' is non-empty.
"""
from __future__ import annotations
from pathlib import Path
from typing import Optional

from docx import Document

from ..store import PlaybookStore

# Column header needles (prefix-match to tolerate trailing whitespace / minor variation)
HEADERS = ("Topic", "Sources", "AI Playbook Rule", "Desirable Sample Clause", "Undesirable Sample Clause")


def _cell(row, idx: int) -> str:
    """Return stripped text of a table cell, or empty string if out of range."""
    if idx < len(row.cells):
        return row.cells[idx].text.strip()
    return ""


def _find_table(doc: Document) -> Optional[tuple[int, dict]]:
    """Return (table_index, col_map) for the first table whose header matches HEADERS."""
    for ti, table in enumerate(doc.tables):
        if not table.rows:
            continue
        header = [cell.text.strip() for cell in table.rows[0].cells]
        if all(any(h.startswith(needle) for h in header) for needle in HEADERS):
            col_map = {}
            for needle in HEADERS:
                for ci, h in enumerate(header):
                    if h.startswith(needle):
                        col_map[needle] = ci
                        break
            return ti, col_map
    return None


def import_docx(store: PlaybookStore, path: str, *,
                name: Optional[str] = None,
                owner_org: str = "") -> str:
    """Import a desirable/undesirable playbook docx into *store*.

    Returns the playbook_id.
    """
    doc = Document(path)
    result = _find_table(doc)
    if result is None:
        raise ValueError(f"No matching table found in {path!r}. "
                         f"Expected columns: {HEADERS}")
    ti, col_map = result

    pid = store.create_playbook(
        name=name or Path(path).stem,
        owner_org=owner_org,
        source_file=path,
    )

    table = doc.tables[ti]
    n_rules = 0
    for ri, row in enumerate(table.rows[1:], start=1):  # ri is 1-based data row index
        topic = _cell(row, col_map["Topic"])
        sources = _cell(row, col_map["Sources"])
        rule_text = _cell(row, col_map["AI Playbook Rule"])
        desirable = _cell(row, col_map["Desirable Sample Clause"])
        undesirable = _cell(row, col_map["Undesirable Sample Clause"])

        if not rule_text:
            continue

        # Desirable → reference_text (the preferred/model clause language)
        # Undesirable → walkaway_language (the language to avoid / triggers escalation)
        # When Undesirable cell is empty, walkaway_language is None — do NOT
        # fall back to the desirable clause; that would corrupt the alignment
        # engine's anti-pattern matching (reference_text ≠ walkaway_language).
        reference_text = desirable or None
        walkaway_language = undesirable or None

        rid = store.create_rule(
            playbook_id=pid,
            title=topic or f"Rule {ri}",
            applies_to="cluster",
            severity="warn",
            nl_assertion=rule_text,
            reference_text=reference_text,
            walkaway_language=walkaway_language,
            rationale=sources or None,
            similarity_threshold=0.85,
            source_provenance={"file": path, "table": ti, "row": ri},
        )
        store.add_binding(
            rule_id=rid,
            entity_kind="cluster",
            entity_id=f"label:{topic}",
            label_text=topic,
            confidence=0.5,
        )
        n_rules += 1

    store.conn.execute(
        "UPDATE playbooks SET description=? WHERE playbook_id=?",
        [f"Imported {n_rules} rules from {Path(path).name}", pid],
    )
    return pid
